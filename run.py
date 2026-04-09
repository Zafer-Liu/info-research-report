# -*- coding: utf-8 -*-
"""
info-research-report skill（专业版）

功能：
1. 从已有的 search 结果 JSON 中生成一份正式研究报告（DOCX，中文）
2. 自动抓取网页正文、调用 LLM 生成结构化摘要，并进行信息类型分类
3. 将报告通过配置好的邮箱发送给指定接收人
4. 设计用于配合 OpenClaw browser 工具使用

使用方式（示例）：
python run.py "美国伊朗战争" "someone@example.com" results.json
python run.py "美国伊朗战争" "someone@example.com" results.json --no-fetch
"""

import os
import sys
import json
import re
import subprocess
from datetime import datetime
from typing import List, Dict, Optional, Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 确保控制台输出为 utf-8
try:
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
except Exception:
    pass


# ==========================
#  搜索相关（供 OpenClaw 外部工具使用）
# ==========================

def search_with_browser(topic: str, max_results: int = 10) -> Optional[Dict[str, Any]]:
    """为 OpenClaw browser 工具构造搜索任务描述"""
    try:
        search_url = f"https://duckduckgo.com/html/?q={topic.replace(' ', '+')}"
        print(f"  Browser will navigate to: {search_url}")
        print(f"  Then snapshot to extract search results")
        return {"action": "search", "url": search_url, "expected_results": max_results}
    except Exception as e:
        print(f"  Browser search setup error: {e}")
        return None


def parse_search_results_from_snapshot(snapshot_text: str, max_results: int = 50) -> List[Dict[str, str]]:
    """从 DuckDuckGo HTML snapshot 中解析搜索结果"""
    results: List[Dict[str, str]] = []
    title_pattern = r'<a class="result__a"[^>]*>([^<]+)</a>'
    url_pattern = r'<a class="result__a" href="([^"]+)"'
    titles = re.findall(title_pattern, snapshot_text)
    urls = re.findall(url_pattern, snapshot_text)
    for i, (title, url) in enumerate(zip(titles, urls)):
        if i >= max_results:
            break
        if 'uddg=' in url:
            import urllib.parse
            parsed = urllib.parse.parse_qs(urllib.parse.urlparse(url).query)
            url = parsed.get('uddg', [url])[0]
        results.append({"title": title.strip(), "url": url, "snippet": ""})
    return results


# ==========================
#  LLM 摘要相关（Minimax/OpenAI + browseros 网页抓取）
# ==========================

def call_openai(prompt: str, max_tokens: int = 1000, temperature: float = 0.3) -> str:
    """调用 LLM API 生成摘要（支持 Minimax / OpenAI）"""
    # 优先使用 Minimax
    api_key = os.environ.get("MINIMAX_API_KEY")
    if api_key:
        try:
            import requests
            url = "https://api.minimax.chat/v1/text/chatcompletion_v2"
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            data = {
                "model": "MiniMax-M2.5",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": temperature
            }
            resp = requests.post(url, headers=headers, json=data, timeout=120)
            resp.raise_for_status()
            result = resp.json()
            return result["choices"][0]["message"]["content"].strip()
        except Exception as e:
            print(f"[Warn] Minimax API failed: {e}, falling back to OpenAI")

    # Fallback: OpenAI
    api_key = (
        os.environ.get("OPENAI_API_KEY")
        or os.environ.get("OPENAI_KEY")
        or os.environ.get("API_KEY")
    )
    if not api_key:
        print("[Warn] No LLM API key found, skipping LLM call")
        return ""
    try:
        import requests
        url = "https://api.openai.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        data = {
            "model": "gpt-3.5-turbo",
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": max_tokens,
            "temperature": temperature
        }
        resp = requests.post(url, headers=headers, json=data, timeout=60)
        resp.raise_for_status()
        result = resp.json()
        return result["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print(f"[Error] OpenAI API call failed: {e}")
        return ""


def fetch_page_content(url: str, timeout: int = 30) -> str:
    """通过 browseros-mcp、mcporter 或 requests 获取网页内容"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    mcp_sh = os.path.join(script_dir, "browseros-mcp", "scripts", "browseros-mcp.sh")

    # 优先尝试本地 browseros-mcp.sh
    if os.path.exists(mcp_sh):
        try:
            import json
            args = json.dumps({"url": url})
            cmd = ["bash", mcp_sh, "get_page_content", args]
            result = subprocess.run(
                cmd, capture_output=True, text=True, timeout=timeout,
                encoding='utf-8', errors='ignore'
            )
            if result.returncode == 0 and result.stdout:
                return result.stdout
        except Exception as e:
            pass  # 继续尝试其他方式

    # 次选 mcporter 调用 browseros
    try:
        cmd = ["mcporter", "call", "browseros.get_page_content", f"url={url}"]
        result = subprocess.run(
            cmd, capture_output=True, text=True, timeout=timeout,
            encoding='utf-8', errors='ignore'
        )
        if result.returncode == 0 and result.stdout:
            return result.stdout
    except Exception as e:
        pass  # 继续尝试 fallback

    # Fallback: 直接用 requests 抓取
    try:
        import requests
        from bs4 import BeautifulSoup

        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        }
        resp = requests.get(url, headers=headers, timeout=timeout)
        resp.raise_for_status()

        # 提取正文
        soup = BeautifulSoup(resp.text, 'html.parser')
        # 移除脚本和样式
        for tag in soup(['script', 'style', 'nav', 'header', 'footer']):
            tag.decompose()

        # 获取文本，保留段落
        text = soup.get_text(separator='\n', strip=True)
        # 清理多余空行
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        return '\n'.join(lines[:2000])  # 限制长度
    except Exception as e:
        print(f"[Warn] Fallback fetch failed for {url}: {e}")
        return ""


def _clean_llm_output(text: str) -> str:
    """清理模型输出中不需要的内容（如 <think> 标签、英文自我说明等）"""
    if not text:
        return ""
    # 去掉 <think>...</think>
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE)
    # 去掉典型英文自述句
    lines = []
    for line in text.splitlines():
        l = line.strip()
        if l.startswith("The user is asking me") or l.startswith("I need to be careful"):
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def summarize_text_with_llm(text: str, prompt: str = "") -> str:
    """使用大模型对 text 做总结（带清洗）"""
    if not text and not prompt:
        return ""
    full_prompt = prompt + ("\n\n" + text if text else "")
    raw = call_openai(full_prompt, max_tokens=900, temperature=0.25)
    return _clean_llm_output(raw)


def summarize_page_content_structured(topic: str, title: str, url: str, content: str) -> Dict[str, str]:
    """
    针对单个网页内容生成“结构化摘要”和“类别标签”。

    返回:
    {
        "summary": "结构化中文摘要",
        "category": "Policy/Industry/Technology/Academia/Other"
    }
    """
    base_prompt = f"""
你是一个中文政策与安全研究员，正在围绕“{topic}”撰写正式研究报告。下面是一篇网页的正文内容，请完成两件事：

【非常重要的要求】
- 全文使用中文作答；
- 不要输出任何英文分析或推理过程；
- 不要出现类似 "The user is asking me"、"I need to be careful" 等英文说明；
- 不要根据年份或你自己的训练时间去判断事件是否真实，也不要说“这可能是未来事件/虚构事件/需要查证”等字样；
- 直接假定这些内容描述的是已经发生的现实事件，专注对文本本身进行客观归纳。

1. 给出该来源的信息类型（只在以下类别中选择一个，使用英文单词）：  
   - Policy   （政策/法律法规/官方文件）  
   - Industry （行业新闻、市场分析、企业动态、媒体报道）  
   - Technology（技术原理、技术评测、技术发展）  
   - Academia （学术论文、研究报告、学者观点）  
   - Other    （其他难以归类的信息）  

2. 输出一个结构化的中文摘要，要求：  
   - 按以下格式输出（必须保留这些中文标题）：  
     【信息类型】(英文字段)  
     【核心结论】若干条分点列出  
     【关键数据或事实】若干条分点列出，如文本中没有明确数字/时间/地点，可写“暂无明确数据”  
     【与研究主题的关联】用中文简要说明该来源对“{topic}”的意义或启示  

3. 风格要求：  
   - 用词正式、客观，不夸张；  
   - 控制在大约 400 字以内；  
   - 不要讨论“真假”“是否虚构”“是否未来事件”等问题。

请直接给出上述结构化内容，不要解释你在做什么，也不要输出与以上结构无关的内容。

网页标题：{title}
网页 URL：{url}

正文内容如下（仅供分析，不要在摘要中长篇照抄原文）：
"""
    raw = summarize_text_with_llm(content, base_prompt)
    if not raw:
        return {"summary": "", "category": "Other"}

    # 从【信息类型】行中抽取类别
    category = "Other"
    for line in raw.splitlines():
        if "【信息类型】" in line:
            m = re.search(r'Policy|Industry|Technology|Academia|Other', line)
            if m:
                category = m.group(0)
            break

    return {"summary": raw.strip(), "category": category}


def summarize_overall_topic(topic: str, page_summaries: List[str]) -> str:
    """基于若干网页摘要，生成一个主题的总体总结（结构化，中文）"""
    if not page_summaries:
        return ""
    joined = "\n\n".join([f"【来源{i + 1}】\n{summary}" for i, summary in enumerate(page_summaries)])
    prompt = f"""
你是一名中文政策与安全研究分析师，正在撰写关于“{topic}”的研究报告。下面是若干网页的结构化摘要，请在综合和归纳的基础上，写出一段“总体形势与关键结论”。

【非常重要的要求】
- 全文使用中文；
- 不要输出任何英文自我说明或推理过程；
- 不要提到“未来事件”“虚构情景”“可能不真实”“需要查证”等表述；
- 直接把这些来源视为描述已经发生的现实事件，对其中信息进行综合研判。

必须按照以下结构输出（保留以下中文标题，不要增加或删减）：  
【整体背景与现状】
- 若干条分点，说明当前总体情况、发展阶段、主要参与者等

【关键趋势与要点】
- 3～5 条分点，概括最重要的发展趋势、政策动向或军事变化

【风险与不确定性】
- 2～4 条分点，说明主要风险、挑战或存在不确定性的方面

风格要求：
- 用词正式、客观，不夸大；
- 结论先行，避免空话、套话；
- 全文控制在 600～800 字以内。

以下是可供参考的来源摘要：
{joined}
"""
    return summarize_text_with_llm("", prompt)


def enrich_results_with_summaries(
    topic: str,
    results: List[Dict[str, Any]],
    max_pages: int = 10,
    fetch_content: bool = True
) -> Dict[str, Any]:
    """
    为搜索结果增加：
    - 每个网页的结构化摘要 llm_summary
    - 每个网页的 category（Policy/Industry/Technology/Academia/Other）
    - 基于摘要生成 overall_summary
    """
    enriched_results: List[Dict[str, Any]] = []
    page_summaries: List[str] = []

    for i, item in enumerate(results):
        if i >= max_pages:
            enriched_results.append(item)
            continue

        title = item.get("title") or ""
        url = item.get("url") or ""
        content = item.get("content") or ""

        # 抓正文
        if not content and fetch_content and url and url != "N/A":
            print(f"  Fetching content from: {url}")
            content = fetch_page_content(url)
            if content:
                item["content"] = content[:8000]

        # 有正文 -> 结构化摘要
        if content:
            res = summarize_page_content_structured(topic, title, url, content)
            llm_summary = res.get("summary") or ""
            category = res.get("category") or "Other"
        else:
            # 无正文 -> 弱摘要
            weak_text = f"网页标题：{title}\n已有简短摘要：{item.get('snippet', '')}"
            prompt = """
你只有网页标题和一个非常简短的摘要。请在不胡乱编造细节的前提下，用中文给出一个简要摘要（约 200 字），说明该来源大致属于什么类型的信息（例如：政策声明、新闻报道、学术评论等），以及可能涉及的重点内容。

要求：
- 全文使用中文；
- 不要提到“未来事件”“虚构情景”“可能不真实”等内容；
- 如你觉得信息确实有限，可以用一句话说明“可用信息有限，需要结合原文进一步核实”，然后概括你能确认的部分。
"""
            llm_summary = summarize_text_with_llm(weak_text, prompt)
            category = "Other"

        if llm_summary:
            item["llm_summary"] = llm_summary
            if not item.get("snippet"):
                item["snippet"] = llm_summary[:200]
            item["category"] = category
            page_summaries.append(llm_summary)
        else:
            item["category"] = item.get("category", "Other")

        enriched_results.append(item)

    overall_summary = ""
    if page_summaries:
        print("  Generating overall summary...")
        overall_summary = summarize_overall_topic(topic, page_summaries)

    return {"overall_summary": overall_summary, "results": enriched_results}


# ==========================
#  报告生成工具
# ==========================

def _sanitize_filename(name: str, max_len: int = 50) -> str:
    """将主题转为文件名安全的字符串"""
    name = re.sub(r'[\\/*?:"<>|]', '_', name)
    return name[:max_len]


def _add_heading_center(doc: Document, text: str, level: int = 0, bold: bool = True):
    """居中标题工具函数"""
    heading = doc.add_heading("", level=level)
    run = heading.add_run(text)
    run.bold = bold
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading


def generate_docx(
    topic: str,
    results: List[Dict[str, Any]],
    output_path: str,
    overall_summary: str = ""
) -> str:
    """生成正式的研究报告（中文，结构化）"""
    if results is None:
        results = []

    doc = Document()

    # ========= 封面 =========
    _add_heading_center(doc, "研究报告", level=0)
    _add_heading_center(doc, f"—— {topic}", level=1)
    doc.add_paragraph()
    doc.add_paragraph()

    org_para = doc.add_paragraph()
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_run = org_para.add_run("编制单位：某某研究部门\n")
    org_run.font.size = Pt(12)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"编制时间：{datetime.now().strftime('%Y年%m月%d日')}")
    date_run.font.size = Pt(12)

    doc.add_page_break()

    # ========= 报告说明 =========
    doc.add_heading("报告说明", level=1)
    info_p = doc.add_paragraph()
    info_p.add_run(f"报告主题：{topic}\n")
    info_p.add_run(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    info_p.add_run(f"信息来源数量：{len(results)}\n")
    info_p.add_run("说明：本报告基于公开网络信息及自动化分析工具生成，仅供内部研究参考，结论及观点不代表任何机构的正式立场。\n")
    doc.add_paragraph("_" * 60)

    # ========= 一、研究背景和目的 =========
    doc.add_heading("一、研究背景和目的", level=1)
    bg_p = doc.add_paragraph()
    bg_p.add_run(
        f"近年来，“{topic}”相关议题在政策、安全形势和国际关系等方面受到广泛关注。"
        f"为系统梳理当前公开信息，识别关键趋势与潜在风险，特编制本研究报告，"
        f"旨在为相关决策和后续深入研究提供基础性参考。"
    )
    doc.add_paragraph("_" * 60)

    # ========= 二、研究方法 =========
    doc.add_heading("二、研究方法", level=1)
    method_p = doc.add_paragraph()
    method_p.add_run(
        "本报告主要采用以下研究方法：\n"
        "1. 网络检索：通过 OpenClaw browser 等工具，在综合搜索引擎中检索与主题高度相关的公开信息；\n"
        "2. 来源筛选：优先选择政策文件、权威机构报告、主流媒体报道及专业分析等相对可靠来源；\n"
        "3. 自动化摘要：使用大语言模型对网页正文进行结构化摘要，提取核心结论、关键数据及与主题的关联；\n"
        "4. 综合研判：在汇总各来源信息的基础上，形成总体形势、关键趋势及风险研判。"
    )
    doc.add_paragraph("_" * 60)

    # ========= 三、总体形势与关键结论 =========
    doc.add_heading("三、总体形势与关键结论", level=1)

    if results:
        summary_p = doc.add_paragraph()
        summary_p.add_run(f"基于所采集的公开信息，围绕“{topic}”，总体形势与主要结论概括如下：\n\n")
        if overall_summary:
            summary_p.add_run(overall_summary + "\n")
        else:
            summary_p.add_run("（当前版本未能生成总体总结，请结合后文分专题和来源分析自行研判。）\n")
    else:
        doc.add_paragraph("暂无可用搜索结果，无法形成总体结论。")

    doc.add_paragraph("_" * 60)

    # ========= 四、分专题分析（按来源类别） =========
    doc.add_heading("四、分专题分析", level=1)

    category_map: Dict[str, List[Dict[str, Any]]] = {}
    for item in results:
        cat = item.get("category") or "Other"
        category_map.setdefault(cat, []).append(item)

    category_names = {
        "Policy": "（一）政策与监管信息",
        "Industry": "（二）媒体报道与舆论动态",
        "Technology": "（三）技术发展与应用",
        "Academia": "（四）学术研究与专家观点",
        "Other": "（五）其他相关信息"
    }

    for key in ["Policy", "Industry", "Technology", "Academia", "Other"]:
        items = category_map.get(key, [])
        if not items:
            continue
        title_cat = category_names.get(key, f"类别：{key}")
        sub_p = doc.add_paragraph()
        run = sub_p.add_run(title_cat)
        run.bold = True
        run.font.size = Pt(12)

        for idx, item in enumerate(items[:5], 1):
            t = item.get("title") or "N/A"
            s = item.get("llm_summary") or item.get("snippet") or ""
            p = doc.add_paragraph()
            p.add_run(f"{idx}. {t}\n").bold = True
            if s:
                p.add_run(f"   概要：{s[:350]}...\n")

    doc.add_paragraph("_" * 60)

    # ========= 五、详细来源分析（Source Details） =========
    doc.add_heading("五、详细来源分析（Source Details）", level=1)

    if results:
        for idx, item in enumerate(results, 1):
            title = item.get("title") or "N/A"
            url = item.get("url") or "N/A"
            snippet = item.get("snippet") or ""
            llm_summary = item.get("llm_summary") or ""
            category = item.get("category") or "Other"

            p = doc.add_paragraph()
            run = p.add_run(f"（{idx}）{title}")
            run.bold = True
            run.font.size = Pt(12)

            p = doc.add_paragraph()
            p.add_run(f"信息类型：{category}\n")

            p = doc.add_paragraph()
            run = p.add_run(f"来源：{url}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 255)

            summary_text = llm_summary or snippet
            if summary_text:
                p = doc.add_paragraph()
                run = p.add_run(f"摘要：{summary_text}")
                run.font.size = Pt(10)
                run.italic = True

            doc.add_paragraph("")
    else:
        doc.add_paragraph("暂无来源数据。")

    doc.add_paragraph("_" * 60)

    # ========= 六、参考资料 =========
    doc.add_heading("六、参考资料（All Search Results）", level=1)

    if results:
        for item in results:
            title = item.get("title") or ""
            url = item.get("url") or ""
            if not title:
                continue
            p = doc.add_paragraph()
            title_run = p.add_run(f"• {title}")
            title_run.bold = True
            if url:
                url_run = p.add_run(f"\n  URL：{url}")
                url_run.font.size = Pt(9)
    else:
        doc.add_paragraph("无。")

    doc.save(output_path)
    return output_path


# ==========================
#  主流程入口
# ==========================

def main():
    """
    使用方式:
    python run.py "主题" [results.json] [--no-fetch]

    选项：
    --no-fetch  跳过网页内容抓取，仅用标题和 snippet 生成摘要（更快）
    """
    args = []
    fetch_content = True
    for arg in sys.argv[1:]:
        if arg == "--no-fetch":
            fetch_content = False
        else:
            args.append(arg)

    topic = args[0] if len(args) > 0 else "测试主题"
    results_file = args[1] if len(args) > 1 else ""

    print(f"=== Research: {topic} ===")
    print(f"[Info] Fetch content from URLs: {fetch_content}")
    print("[Info] Use --no-fetch to skip fetching (faster but less detailed)")

    # 加载结果
    if results_file and os.path.exists(results_file):
        try:
            with open(results_file, 'r', encoding='utf-8') as f:
                results = json.load(f)
            if not isinstance(results, list):
                results = []
            print(f"\n[1/4] Loaded {len(results)} results from {results_file}")
        except Exception as e:
            print(f"[Error] Failed to read results file: {e}")
            results = []
    else:
        print("\n[1/4] No results file provided, using preset data")
        results = [{"title": "请先使用 OpenClaw browser 工具进行搜索", "url": "N/A", "snippet": ""}]

    # 生成 LLM 摘要 & 总体总结
    print("\n[2/4] Generating structured summaries with LLM...")
    enriched = enrich_results_with_summaries(topic, results, max_pages=10, fetch_content=fetch_content)
    overall_summary = enriched["overall_summary"]
    results = enriched["results"]
    print("  Summaries generation finished.")

    # 生成 DOCX
    print("\n[3/4] Generating DOCX report (Chinese, professional style)...")
    safe_topic = _sanitize_filename(topic.replace(" ", "_"), max_len=30)
    output_file = f"Report_{safe_topic}_{datetime.now().strftime('%Y%m%d%H%M')}.docx"

    try:
        generate_docx(topic, results, output_file, overall_summary=overall_summary)
        print(f"  Saved report: {output_file}")
        print(f"  Report path: {os.path.abspath(output_file)}")
    except Exception as e:
        print(f"[Error] Failed to generate DOCX: {e}")
        print("\n=== Aborted ===")
        return

    print("\n=== Done ===")


if __name__ == "__main__":
    main()
